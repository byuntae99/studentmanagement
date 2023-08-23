from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponseRedirect, HttpResponse
from django.shortcuts import render, redirect
from django.contrib import messages
from app_001.EmailBackEnd import EmailBackEnd
from django.views.decorators.csrf import csrf_exempt
from app_001 .models import *
from .forms import *
from docx import Document

def home(request):
    return render(request, 'index.html')

def loginPage(request):
        return render(request, 'login.html')

def course_reg_confirm(request):
        return render(request,'hod_template/course_registratin.html')

def atendance_list(request):
    staffid=Staff.objects.get(admin=request.user.id)
    students = Students.objects.all()
    context = {
        "students": students
    }
    return render(request,'staff_template/attendance_tabel.html',context)

    
def doLogin(request):
    if request.method != "POST":
        return HttpResponse("<h2>## Error 404 ##</h2>")
    else:
        username=request.POST.get('username')
        password=request.POST.get('password')
        user = EmailBackEnd.authenticate(request,username, password)
        if user != None:
            login(request, user)
            user_type = user.user_type
            if user_type == '1':
                return redirect('admin_home')
                      
            elif user_type == '2':
                # return HttpResponse("Student Login")
                return redirect('student_home')
            elif user_type == '3':
                # return HttpResponse("Staff Login")
                return redirect('staff_home')
            else:
                messages.error(request, "Invalid Login!")
                return redirect('login')
        else:
            messages.error(request, "Invalid Login Credentials!")
            #return HttpResponseRedirect("/")
            return redirect('login')


def logout_user(request):
    logout(request)
    return HttpResponseRedirect('/')



############_____ADMIN______#############



def adminhome(request):
    all_student_count = Students.objects.all().count()
    all_staff_count = Staff.objects.all().count()
    all_student_leaves=LeaveReportStudent.objects.all().count()
    curs_status=Status.objects.all().count()
    if curs_status :
        curs_status="Yes"
    else:
        curs_status="No"
    context={
        "all_student_count": all_student_count,  
        "all_staff_count": all_staff_count,
        "all_student_leaves" : all_student_leaves,
        "course_status" : curs_status,
    }
    return render(request, "hod_template/home_content.html", context)

def staff_home(request):
    staff_obj = Staff.objects.get(admin=request.user.id)
    return render(request, "staff_template/staff_home_template.html")

def addstaff(request):
    form = AddstaffForm()
    context = {
        "form": form
    }
    return render(request, 'hod_template/add_staff_template.html', context)

def addstaffsave(request):
    
    if request.method != "POST":
        messages.error(request, "Invalid Method")
        return redirect('addstaff')
    else:
        form = AddstaffForm(request.POST,request.FILES)
        if form.is_valid():
            first_name = form.cleaned_data['first_name']
            last_name = form.cleaned_data['last_name']
            username = form.cleaned_data['username']
            email = form.cleaned_data['email']
            password = form.cleaned_data['password']
            address = form.cleaned_data['address']
            gender = form.cleaned_data['gender']
            image_file=form.cleaned_data['image_file']
            user = CustomUser.objects.create_user(username=username, password=password, email=email,first_name=first_name, last_name=last_name, user_type=3)
            user.staff.name=first_name +" "+ last_name
            user.staff.email=email
            user.staff.password=password
            user.staff.address = address
            
            user.staff.gender = gender
            user.staff.imagefiled=image_file
        
            user.save()
            
            messages.success(request, "Staff Added Successfully!")
            return redirect('addstaff')
        
        else:
            messages.error(request, "Failed to Add Staff!")
            return redirect('addstaff')

def manage_staff(request):
    staffs = Staff.objects.all()
    context = {
        "staffs": staffs,
    }
    return render(request, 'hod_template/manage_staff_template.html', context)

def edit_staff(request):
    return render(request, 'hod_template/manage_staff_template.html')

def add_course(request):
    staffname=Staff.objects.all()
    context = {
        "staffname": staffname
    }
    return render(request,'hod_template/add_course.html',context)

def add_course_save(request):
    if request.method != "POST":
        messages.error(request, "Invalid Method")
        return redirect('add_course')
    else:
        if request.method == "POST":
            stfname=request.POST["staffname"]
            curname=request.POST["coursename"]
            sa=Addcourse.objects.create(staff_name=stfname)
            sa.save()
            messages.success(request, "Course Added Successfully!")
            return redirect('add_course')
        else:
            messages.error(request, "Failed to Add Course!")
            return redirect('add_course')
        
def addstudent(request):
    form = AddStudentForm()
    context = {
        "form": form
    }
    return render(request, 'hod_template/add_student_template.html',context)


def addstudentsave(request):
    if request.method != "POST":
        messages.error(request, "Invalid Method")
        return redirect('add_student')
    else:
        form = AddStudentForm(request.POST,request.FILES)
        if form.is_valid():
            first_name = form.cleaned_data['first_name']
            last_name = form.cleaned_data['last_name']
            username = form.cleaned_data['username']
            email = form.cleaned_data['email']
            rollno=form.cleaned_data['rollnumber']
            password = form.cleaned_data['password']
            address = form.cleaned_data['address']
            gender = form.cleaned_data['gender']
            image_file=form.cleaned_data['image_file']
            user = CustomUser.objects.create_user(username=username, password=password, email=email,first_name=first_name, last_name=last_name, user_type=2)
            user.students.name=first_name +' '+ last_name
            user.students.email=email
            user.students.password=password
            user.students.address = address
            user.students.rollnumber = rollno
            user.students.gender = gender
            user.students.imagefiled=image_file
        
            user.save()
            
            messages.success(request, "Student Added Successfully!")
            return redirect('add_student')
        
        else:
            messages.error(request, "Failed to Add Student!")
            return redirect('add_student')


def managestudent(request):
    students = Students.objects.all()
    context = {
        "students": students
    }
    return render(request, 'hod_template/manage_student_template.html', context)


def editstudent(request, student_id):
    # Adding Student ID into Session Variable
    request.session['student_id'] = student_id
    student = Students.objects.get(admin=student_id)
    form = EditStudentForm()
    # Filling the form with Data from Database
    form.fields['email'].initial = student.admin.email
    form.fields['username'].initial = student.admin.username
    form.fields['rollnumber'].initial = student.rollnumber
    form.fields['first_name'].initial = student.admin.first_name
    form.fields['last_name'].initial = student.admin.last_name
    form.fields['address'].initial = student.address
    form.fields['gender'].initial = student.gender
    form.fields['image_file'].initial=student.imagefiled
    context = {
        "id": student_id,
        "username": student.admin.username,
        "form": form
    }
    return render(request, "hod_template/edit_student_template.html", context)


def edit_student_save(request): 

    if request.method != "POST":
        return HttpResponse("Invalid Method!")
    else:
        student_id = request.session.get('student_id')
        #print("@@@@@@@@@@@@@@@@@@@@@@@#################$%^&*(*&^%$#$%^&*&^%$%^^^^^^^^^^)",student_id)
        if student_id == None:
            print("abcdefghijklmnopqrstuvxyz")
            return redirect('/manage_student')

        form = EditStudentForm(request.POST,request.FILES)
        email = form.data['email']
        username = form.data['username']
        first_name = form.data['first_name']
        last_name = form.data['last_name']
        address = form.data['address']
        gender = form.data['gender']
        image=form.data['image_file']
        try:
                # First Update into Custom User Model
                usera = CustomUser.objects.get(id=student_id)
               
                usera.first_name = first_name
                usera.last_name = last_name
                usera.email = email
                usera.username = username
                usera.save()
                student_obj=Students.objects.get(admin=student_id)
                student_obj.name=first_name +' '+ last_name
                student_obj.email=email
                student_obj.address = address
                student_obj.gender = gender
                student_obj.imagefiled.clear()
                student_obj.imagefiled=image
                print("!@#$%^&*()(*&^%$#@!@#$%^&*())(*&^%$#@!@#$%^&*()(*&^%$#@@#$%^&*()(*&^%$#@!@#$%^&*()(*&^%$#@#$%$#@#$%^&*()))))))",image)

                student_obj.save()

                del request.session['student_id']

                messages.success(request, "Student Updated Successfully!")
                return redirect('/edit_student/'+f"{student_id}")
        except:
                messages.error(request, "Failed to Update Student.")
                return redirect('/edit_student/'+f'{student_id}')
        
    


def deletestudent(request, student_id):
    student = Students.objects.get(admin=student_id)
    student_email=student.email
    user_obj = CustomUser.objects.get(email=student_email)
    try:
        student.delete()
        user_obj.delete()
        messages.success(request, "Student Deleted Successfully.")
        return redirect('manage_student')
    except:
        messages.error(request, "Failed to Delete Student.")
        return redirect('manage_student')

def course_reg_update(request):
    if request.method == "POST":
        aa=int(request.POST["radio"])
        if aa==1:
            ss=Status.objects.create(status=aa)
        elif aa==0:
            ss=Status.objects.all()
            ss.delete()    
        return render(request,'hod_template/course_registratin.html')
    
    
@csrf_exempt
def checkemailexist(request):
    email = request.POST.get("email")
    user_obj = CustomUser.objects.filter(email=email)
    std=Students.objects.filter(email=email)
    stf=Staff.objects.filter(email=email)
    print('###########################@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@####################',user_obj,std,stf)
    if (user_obj or std or stf) :
        return HttpResponse(True)
    else:
        return HttpResponse(False)

@csrf_exempt
def checkusernameexist(request):
    username = request.POST.get("username")
    user_obj = CustomUser.objects.filter(username=username).exists()
    if user_obj:
        return HttpResponse(True)
    else:
        return HttpResponse(False)
    
@csrf_exempt
def check_Staff_exist(request):
    stf=request.POST.get("staffname")
    user_staff=Addcourse.objects.filter(staff_name=stf).exists()
    if user_staff:
        return HttpResponse(True)
    else:
        return HttpResponse(False)

##______STUDENT______##



def studenthome(request):
    student_obj = Students.objects.get(admin=request.user.id)
    return render(request, "student_template/student_home_template.html")


def studentapplyleave(request):
    student_obj = Students.objects.get(admin=request.user.id)
    leave_data = LeaveReportStudent.objects.filter(student_id=student_obj)
    context = {
        "leave_data": leave_data
    }
    return render(request, 'student_template/student_apply_leave.html', context)


def studentapplyleavesave(request):
    if request.method != "POST":
        messages.error(request, "Invalid Method")
        return redirect('student_apply_leave')
    else:
        leavetype= request.POST['leave_type']
        fromleave_date = request.POST['fromleave_date']
        toleave_date = request.POST['toleave_date']
        leave_message = request.POST['leave_message']

        student_obj = Students.objects.get(admin=request.user.id)
        try:
            leave_report = LeaveReportStudent(student_id=student_obj, leave_types=leavetype, from_leave_date=fromleave_date,to_leave_date=toleave_date,leave_message=leave_message, leave_status=0)
            leave_report.save()
            messages.success(request, "Applied for Leave.")
            return redirect('student_apply_leave')
        except:
            messages.error(request, "Failed to Apply Leave")
            return redirect('student_apply_leave')
        
def cancelleave(request): 
    student_obj = Students.objects.get(admin=request.user.id)
    leave_deletes=LeaveReportStudent.objects.filter(student_id=student_obj)
    leave_deletes.delete()
    return redirect('student_apply_leave')

#########-LEAVE VIEWS-#################
def studentleaveview(request):
    leaves = LeaveReportStudent.objects.all()
    context = {
        "leaves": leaves
    }
    return render(request, 'hod_template/student_leave_view.html', context)

def studentleaveapprove(request, leave_id):
    leave = LeaveReportStudent.objects.get(id=leave_id)
    leave.leave_status = 1
    leave.save()
    return redirect('student_leave_view')


def studentleavereject(request, leave_id):
    leave = LeaveReportStudent.objects.get(id=leave_id)
    leave.leave_status = 2
    leave.save()
    return redirect('student_leave_view')

def viewprofile(request):
    return render(request,'student_template/student_profile.html')


def student_register(request):
    aa=Status.objects.filter(status=1)
    bb=CouRegForm()
    return render(request,'student_template/course_reg_view.html',{'aa' : aa , 'bb':bb})

def crssavwe(request):
    student_obj = Students.objects.get(admin=request.user.id)  
    if request.method =="POST":
        stff=request.POST['stafflist']
        curs=request.POST['corselist']
        tms=request.POST['timess']
        aba=Courseregister.objects.create(Stafff=stff,Coursee=curs,timess=tms,studentid=student_obj)
        messages.success(request,'Course Registered Sucessfuly')
    return redirect('student_register')

def dowload_course_doc(request):
    document = Document()
    sampless=Students.objects.get(admin=request.user.id)
    document.add_heading('Course Registered', 0)
    p = document.add_paragraph('Time Table ')
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Course'
    hdr_cells[1].text = 'Staff'
    hdr_cells[2].text = 'Time'
    curs=Courseregister.objects.filter(studentid=sampless)
    for i in curs:    
        row = table.add_row().cells
        row[0].text = i.Coursee
        row[1].text = i.Stafff
        row[2].text = i.timess
    #document.add_page_break()
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{sampless.rollnumber}".docx'
    document.save(response)
    return response
 

